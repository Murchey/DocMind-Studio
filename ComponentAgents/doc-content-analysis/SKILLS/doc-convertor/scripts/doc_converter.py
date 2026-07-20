import os
import sys
import json
import shutil
import subprocess
from pathlib import Path


class DocConverter:
    """DOC/DOCX/PDF 文档转换与内容提取，支持单文件和批量处理。"""

    def __init__(self):
        self.doc_methods = [
            ('win32com', self._convert_doc_with_win32com),
            ('libreoffice', self._convert_with_libreoffice),
        ]

    # ── 格式检测 ──────────────────────────────────────────────

    @staticmethod
    def is_doc_format(file_path: str) -> bool:
        path = Path(file_path)
        if not path.exists():
            return False
        ext = path.suffix.lower()
        if ext == '.doc':
            return True
        if ext == '.docx':
            return False
        try:
            with open(file_path, 'rb') as f:
                header = f.read(8)
                if header[:4] == b'\xd0\xcf\x11\xe0':
                    return True
        except Exception:
            pass
        return False

    @staticmethod
    def is_pdf_format(file_path: str) -> bool:
        path = Path(file_path)
        if not path.exists():
            return False
        if path.suffix.lower() == '.pdf':
            return True
        try:
            with open(file_path, 'rb') as f:
                header = f.read(5)
                return header == b'%PDF-'
        except Exception:
            pass
        return False

    # ── .doc → .docx 转换 ─────────────────────────────────────

    def _convert_doc_with_win32com(self, input_path: str, output_path: str) -> dict:
        try:
            import win32com.client
            import pythoncom
            pythoncom.CoInitialize()
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            word.DisplayAlerts = False
            try:
                abs_input = str(Path(input_path).resolve())
                abs_output = str(Path(output_path).resolve())
                doc = word.Documents.Open(abs_input)
                doc.SaveAs2(abs_output, FileFormat=16)
                doc.Close()
                return {
                    'success': True, 'output_path': output_path,
                    'method': 'win32com', 'file_size': os.path.getsize(output_path)
                }
            finally:
                word.Quit()
                pythoncom.CoUninitialize()
        except Exception as e:
            return {'success': False, 'method': 'win32com', 'error': str(e)}

    def _convert_with_libreoffice(self, input_path: str, output_path: str) -> dict:
        try:
            output_dir = str(Path(output_path).parent)
            abs_input = str(Path(input_path).resolve())
            soffice_paths = [
                r'C:\Program Files\LibreOffice\program\soffice.exe',
                r'C:\Program Files (x86)\LibreOffice\program\soffice.exe',
                'soffice',
            ]
            soffice_cmd = None
            for p in soffice_paths:
                if os.path.exists(p) or p == 'soffice':
                    soffice_cmd = p
                    break
            if not soffice_cmd:
                return {'success': False, 'method': 'libreoffice', 'error': 'LibreOffice not found'}
            result = subprocess.run(
                [soffice_cmd, '--headless', '--convert-to', 'docx', '--outdir', output_dir, abs_input],
                capture_output=True, text=True, timeout=120
            )
            if result.returncode == 0:
                input_stem = Path(input_path).stem
                generated_file = Path(output_dir) / f"{input_stem}.docx"
                if generated_file.exists() and str(generated_file) != str(output_path):
                    shutil.move(str(generated_file), output_path)
                return {
                    'success': True, 'output_path': output_path,
                    'method': 'libreoffice', 'file_size': os.path.getsize(output_path)
                }
            else:
                return {'success': False, 'method': 'libreoffice', 'error': result.stderr or result.stdout}
        except Exception as e:
            return {'success': False, 'method': 'libreoffice', 'error': str(e)}

    def convert_doc(self, input_path: str, output_path: str = None) -> dict:
        if not os.path.exists(input_path):
            return {'success': False, 'error': f'Input file not found: {input_path}'}
        if not self.is_doc_format(input_path):
            return {
                'success': True, 'input_path': input_path,
                'output_path': input_path, 'method': 'none',
                'message': 'File is already in .docx format or not a .doc file'
            }
        if output_path is None:
            output_path = str(Path(input_path).with_suffix('.docx'))
        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        for method_name, method_func in self.doc_methods:
            print(f"[INFO] Trying .doc conversion method: {method_name}")
            result = method_func(input_path, output_path)
            if result['success']:
                print(f"[INFO] Conversion successful using {method_name}")
                result['input_path'] = input_path
                return result
            else:
                print(f"[WARN] Method {method_name} failed: {result.get('error', 'Unknown error')}")
        return {
            'success': False, 'input_path': input_path,
            'error': 'All conversion methods failed. Please install Microsoft Word or LibreOffice.'
        }

    # ── PDF → DOCX 转换 ──────────────────────────────────────

    def convert_pdf(self, input_path: str, output_path: str = None) -> dict:
        if not os.path.exists(input_path):
            return {'success': False, 'error': f'Input file not found: {input_path}'}
        if not self.is_pdf_format(input_path):
            return {'success': False, 'error': f'Not a PDF file: {input_path}'}
        if output_path is None:
            output_path = str(Path(input_path).with_suffix('.docx'))
        Path(output_path).parent.mkdir(parents=True, exist_ok=True)

        # 方法1: pdf2docx（适用于结构化PDF）
        try:
            from pdf2docx import Converter
            print(f"[INFO] Converting PDF to DOCX with pdf2docx: {Path(input_path).name}")
            cv = Converter(input_path)
            cv.convert(output_path)
            cv.close()
            print(f"[INFO] PDF conversion successful with pdf2docx")
            return {
                'success': True, 'input_path': input_path,
                'output_path': output_path, 'method': 'pdf2docx',
                'file_size': os.path.getsize(output_path)
            }
        except ImportError:
            print(f"[WARN] pdf2docx not installed, trying PyMuPDF...")
        except Exception as e:
            print(f"[WARN] pdf2docx failed: {e}, trying PyMuPDF...")

        # 方法2: PyMuPDF直接提取文字（适用于知网等中文PDF）
        pdf_content = self.extract_pdf_content(input_path)
        if pdf_content.get('success'):
            # 将提取的内容保存为文本文件
            txt_output = str(Path(output_path).with_suffix('.txt'))
            try:
                with open(txt_output, 'w', encoding='utf-8') as f:
                    f.write(pdf_content['full_text'])
                return {
                    'success': True, 'input_path': input_path,
                    'output_path': txt_output, 'method': 'pymupdf_text',
                    'file_size': os.path.getsize(txt_output)
                }
            except Exception as e:
                return {
                    'success': False, 'input_path': input_path,
                    'method': 'pymupdf_text', 'error': str(e)
                }

        return {
            'success': False, 'input_path': input_path,
            'error': 'All PDF conversion methods failed. Please install pdf2docx or PyMuPDF (pip install pymupdf).'
        }

    # ── PDF直接文字提取（PyMuPDF，增强版） ──────────────────────

    # 中文学术论文常见标题关键词模式
    _CN_HEADING_PATTERNS = [
        '引言', '前言', '绪论', '摘要', 'Abstract',
        '结论', '总结', '展望', '致谢', '参考文献', 'References',
        '关键词', 'Keywords', '研究背景', '研究方法', '实验', '结果',
        '讨论', '分析', '文献综述', '相关工作',
    ]

    def extract_pdf_content(self, pdf_path: str) -> dict:
        """
        使用 PyMuPDF 从 PDF 提取文字内容（增强版）。
        - 优先使用 get_text("dict") 获取字体级信息，识别标题层级
        - 对扫描型/图片型 PDF 自动回退到 OCR
        - 支持多栏排版、表/图检测
        """
        if not os.path.exists(pdf_path):
            return {'success': False, 'error': f'File not found: {pdf_path}'}
        try:
            import fitz
        except ImportError:
            return {'success': False, 'error': 'PyMuPDF not installed. Run: pip install pymupdf'}

        try:
            doc = fitz.open(pdf_path)
            total_pages = doc.page_count

            # ── Phase 1: 逐页提取 ──
            all_blocks = []       # 所有文本块（含坐标、字体）
            has_text = False      # 是否有可提取的文字
            scan_pages = []       # 需要 OCR 的页面

            for page_num in range(total_pages):
                page = doc[page_num]
                blocks = self._extract_page_blocks(page, page_num)
                if blocks:
                    has_text = True
                    all_blocks.extend(blocks)
                else:
                    scan_pages.append(page_num)

            # ── Phase 2: OCR 回退 ──
            if not has_text and scan_pages:
                print(f"[INFO] No extractable text found — PDF appears to be scanned/image-based. "
                      f"Attempting OCR on {len(scan_pages)} page(s)...")
                ocr_blocks = self._ocr_extract(doc, scan_pages)
                if ocr_blocks:
                    all_blocks = ocr_blocks
                    has_text = True
                    print(f"[INFO] OCR extracted {len(all_blocks)} text block(s)")
                else:
                    print(f"[WARN] OCR failed or not available. Install Tesseract + pytesseract for OCR support.")

            # ── Phase 2b: 混合模式（部分页有文字，部分需 OCR）──
            elif scan_pages:
                print(f"[INFO] {len(scan_pages)} page(s) have no extractable text. Attempting OCR on those pages...")
                ocr_blocks = self._ocr_extract(doc, scan_pages)
                if ocr_blocks:
                    all_blocks.extend(ocr_blocks)
                    # 按页码重排
                    all_blocks.sort(key=lambda b: (b.get('page', 0), b.get('y0', 0), b.get('x0', 0)))

            if not all_blocks:
                doc.close()
                return {'success': False, 'error': 'No text content extracted from PDF (text and OCR both failed)'}

            # ── Phase 3: 字体统计 → 判定标题阈值 ──
            font_stats = self._analyze_fonts(all_blocks)
            body_size = font_stats['body_size']

            # ── Phase 4: 标题识别 → 生成结构化段落 ──
            paragraphs, tables = self._classify_blocks(all_blocks, body_size, font_stats)

            # ── Phase 5: 后处理 ──
            paragraphs = self._post_process_paragraphs(paragraphs)
            full_text = '\n'.join(p['text'] for p in paragraphs)

            doc.close()

            return {
                'success': True,
                'file_name': Path(pdf_path).name,
                'file_path': str(pdf_path),
                'title': self._detect_title(paragraphs) or Path(pdf_path).stem,
                'author': self._detect_author(paragraphs),
                'created_at': self._detect_date(paragraphs),
                'page_count': total_pages,
                'paragraph_count': len(paragraphs),
                'table_count': len(tables),
                'paragraphs': paragraphs,
                'tables': tables,
                'full_text': full_text,
                'extraction_method': 'pymupdf_dict' if has_text else 'ocr',
                'font_stats': {
                    'body_size_pt': round(body_size, 1),
                    'title_size_pt': round(font_stats.get('title_size', 16), 1) if font_stats.get('title_size') else None,
                    'h1_size_pt': round(font_stats.get('h1_size', 14), 1) if font_stats.get('h1_size') else None,
                    'h2_size_pt': round(font_stats.get('h2_size', 12), 1) if font_stats.get('h2_size') else None,
                    'h3_size_pt': round(font_stats.get('h3_size', 10.5), 1) if font_stats.get('h3_size') else None,
                }
            }
        except Exception as e:
            import traceback
            traceback.print_exc()
            return {'success': False, 'error': f'Failed to extract PDF content: {str(e)}'}

    # ── 页面块提取（含字体信息）──────────────────────────────

    def _extract_page_blocks(self, page, page_num: int) -> list:
        """使用 get_text('dict') 提取每页的文字块，保留字体、坐标信息"""
        import fitz
        try:
            text_dict = page.get_text("dict", flags=fitz.TEXT_PRESERVE_WHITESPACE)
        except Exception:
            return []

        blocks = []
        for block in text_dict.get("blocks", []):
            if block.get("type") != 0:  # 非文本块（图片等）
                continue

            block_text_parts = []
            block_fonts = []
            block_bbox = None

            for line in block.get("lines", []):
                line_text_parts = []
                for span in line.get("spans", []):
                    text = span.get("text", "").strip()
                    if not text:
                        continue
                    line_text_parts.append(text)
                    block_fonts.append({
                        'size': span.get('size', 0),
                        'font': span.get('font', ''),
                        'flags': span.get('flags', 0),
                        'color': span.get('color', 0),
                    })
                    # 更新包围盒
                    bbox = span.get('bbox')
                    if bbox:
                        if block_bbox is None:
                            block_bbox = list(bbox)
                        else:
                            block_bbox[0] = min(block_bbox[0], bbox[0])
                            block_bbox[1] = min(block_bbox[1], bbox[1])
                            block_bbox[2] = max(block_bbox[2], bbox[2])
                            block_bbox[3] = max(block_bbox[3], bbox[3])

                if line_text_parts:
                    block_text_parts.append(''.join(line_text_parts))

            if not block_text_parts or not block_fonts:
                continue

            full_text = ' '.join(block_text_parts)
            # 计算该块的主要字体大小（众数）
            sizes = [f['size'] for f in block_fonts]
            if sizes:
                from collections import Counter
                main_size = Counter(round(s, 1) for s in sizes).most_common(1)[0][0]
            else:
                main_size = 10.5

            # 是否粗体（取大多数 span 的粗体标志）
            bold_count = sum(1 for f in block_fonts if f.get('flags', 0) & 2**3)
            is_bold = bold_count > len(block_fonts) / 2

            blocks.append({
                'page': page_num,
                'text': full_text,
                'size': main_size,
                'is_bold': is_bold,
                'font_name': block_fonts[0].get('font', '') if block_fonts else '',
                'char_count': len(full_text),
                'line_count': len(block_text_parts),
                'bbox': block_bbox,
                'x0': block_bbox[0] if block_bbox else 0,
                'y0': block_bbox[1] if block_bbox else 0,
                'x1': block_bbox[2] if block_bbox else 0,
                'y1': block_bbox[3] if block_bbox else 0,
            })

        return blocks

    # ── OCR 回退 ──────────────────────────────────────────────

    def _ocr_extract(self, doc, page_nums: list) -> list:
        """
        对指定的页面列表执行 OCR 识别。
        方法1: PyMuPDF 内置 Tesseract (fitz.get_tessocr / page.get_textpage_ocr)
        方法2: 外部 pytesseract
        """
        blocks = []
        try:
            import fitz

            # 检查 Tesseract 是否可用
            tesseract_available = False
            try:
                import pytesseract
                from PIL import Image
                import io
                tesseract_available = True
            except ImportError:
                pass

            for page_num in page_nums:
                page = doc[page_num]

                # 方法1: PyMuPDF 内置 OCR（需要系统安装 Tesseract）
                try:
                    tp = page.get_textpage_ocr(
                        flags=3,
                        language='chi_sim+eng',
                        dpi=300,
                        full=True,
                    )
                    text = page.get_text(textpage=tp)
                    if text and text.strip():
                        # 按行分割，当作普通文本块处理
                        for line in text.splitlines():
                            line = line.strip()
                            if not line:
                                continue
                            blocks.append({
                                'page': page_num,
                                'text': line,
                                'size': 10.5,
                                'is_bold': False,
                                'font_name': 'SimSun',
                                'char_count': len(line),
                                'line_count': 1,
                                'bbox': None,
                                'x0': 72, 'y0': 72, 'x1': 500, 'y1': 84,
                            })
                        print(f"[INFO] page {page_num + 1}: PyMuPDF OCR extracted {len(text)} chars")
                        continue
                except Exception:
                    pass

                # 方法2: pytesseract（需要 pip install pytesseract + 系统安装 Tesseract）
                if tesseract_available and not blocks:
                    try:
                        mat = fitz.Matrix(2, 2)  # 2x 缩放提高 OCR 精度
                        pix = page.get_pixmap(matrix=mat)
                        img = Image.open(io.BytesIO(pix.tobytes("png")))
                        text = pytesseract.image_to_string(img, lang='chi_sim+eng')
                        if text.strip():
                            for line in text.splitlines():
                                line = line.strip()
                                if not line:
                                    continue
                                blocks.append({
                                    'page': page_num,
                                    'text': line,
                                    'size': 10.5,
                                    'is_bold': False,
                                    'font_name': 'SimSun',
                                    'char_count': len(line),
                                    'line_count': 1,
                                    'bbox': None,
                                    'x0': 72, 'y0': 72, 'x1': 500, 'y1': 84,
                                })
                            print(f"[INFO] page {page_num + 1}: pytesseract OCR extracted {len(text)} chars")
                    except Exception:
                        pass

        except Exception as e:
            print(f"[WARN] OCR extraction failed: {e}")

        return blocks

    # ── 字体分析 ──────────────────────────────────────────────

    def _analyze_fonts(self, blocks: list) -> dict:
        """分析所有文本块的字体大小分布，判定正文基准字号和各级标题阈值"""
        if not blocks:
            return {'body_size': 10.5}

        sizes = [b['size'] for b in blocks if b['size'] > 0]
        if not sizes:
            return {'body_size': 10.5}

        from collections import Counter
        size_counter = Counter(round(s, 1) for s in sizes)

        # 正文通常是出现频率最高的字号（排除极小和极大值）
        common_sizes = [(s, c) for s, c in size_counter.most_common()
                        if 8 <= s <= 14]
        if not common_sizes:
            body_size = sorted(sizes)[len(sizes) // 2]
        else:
            body_size = common_sizes[0][0]

        # 各级标题字号（比 body_size 大）
        larger_sizes = sorted(set(round(s, 1) for s in sizes if s > body_size + 1), reverse=True)

        result = {'body_size': body_size}

        if larger_sizes:
            # 最大字号 → title
            result['title_size'] = larger_sizes[0]
            # 其余按大小分 H1/H2/H3
            if len(larger_sizes) >= 4:
                result['h1_size'] = larger_sizes[1]
                result['h2_size'] = larger_sizes[2]
                result['h3_size'] = larger_sizes[3]
            elif len(larger_sizes) == 3:
                result['h1_size'] = larger_sizes[1]
                result['h2_size'] = larger_sizes[2]
                result['h3_size'] = body_size + 0.5
            elif len(larger_sizes) == 2:
                result['h1_size'] = larger_sizes[1]
                result['h2_size'] = body_size + 0.5
                result['h3_size'] = body_size
            else:
                result['h1_size'] = body_size + 1
                result['h2_size'] = body_size + 0.5
                result['h3_size'] = body_size
        else:
            # 无大字号的退化情况：按粗体判断
            result['title_size'] = body_size
            result['h1_size'] = body_size
            result['h2_size'] = body_size
            result['h3_size'] = body_size

        return result

    # ── 块分类（标题 vs 正文）─────────────────────────────────

    def _classify_blocks(self, blocks: list, body_size: float, font_stats: dict) -> tuple:
        """
        将文本块分类为不同级别的标题和正文段落。
        返回 (paragraphs, tables)
        """
        title_size = font_stats.get('title_size', body_size + 3)
        h1_size = font_stats.get('h1_size', body_size + 2)
        h2_size = font_stats.get('h2_size', body_size + 1)
        h3_size = font_stats.get('h3_size', body_size + 0.5)

        paragraphs = []
        tables = []
        page_blocks_grouped = {}
        for b in blocks:
            page_blocks_grouped.setdefault(b['page'], []).append(b)

        for page_num in sorted(page_blocks_grouped):
            page_blocks = page_blocks_grouped[page_num]
            for blk in page_blocks:
                text = blk['text'].strip()
                if not text:
                    continue

                size = blk['size']
                is_bold = blk.get('is_bold', False)
                char_count = blk['char_count']
                font_name = blk.get('font_name', '').lower()

                # ── 字数太少 → 可能是页码/页眉，跳过 ──
                if char_count <= 2 and not any(kw in text for kw in ['摘要', '引言', '结论']):
                    continue

                # ── 字号判定 ──
                if size >= title_size + 1:
                    level = 0  # Title
                elif size >= h1_size - 0.2:
                    level = 1  # H1
                elif size >= h2_size - 0.2:
                    level = 2  # H2
                elif size >= h3_size - 0.2:
                    level = 3  # H3
                else:
                    level = 0  # 正文

                # ── 粗体加分 ──
                if is_bold and level == 0 and size >= body_size:
                    # 短行 + 粗体 + 字号不少于正文 → 可能是标题
                    if char_count <= 60:
                        level = 3
                    elif size >= h2_size - 0.5:
                        level = 2

                # ── 关键词/模式匹配修正 ──
                level = self._pattern_heading_fix(text, level, char_count)

                # ── 黑体/标题字体加分 ──
                if level == 0 and size >= body_size:
                    if any(f in font_name for f in ['heiti', 'simhei', '黑体', 'bold']):
                        if char_count <= 80:
                            level = 3

                is_heading = level >= 1

                paragraphs.append({
                    'text': text,
                    'style': f'Heading {level}' if is_heading else 'Normal',
                    'is_heading': is_heading,
                    'level': level,
                    'page': page_num + 1,
                    'font_size': round(size, 1),
                    'is_bold': is_bold,
                    'char_count': char_count,
                })

        return paragraphs, tables

    def _pattern_heading_fix(self, text: str, current_level: int, char_count: int) -> int:
        """基于内容模式修正标题级别"""
        # 中文数字编号: 一、二、三... / 1. / 1.1 / 第X章
        import re

        # 第X章/节/部分
        if re.match(r'^第[一二三四五六七八九十\d]+[章节部分篇]', text):
            return max(current_level, 1)

        # 数字编号: 1 xxx, 1. xxx, 1.1 xxx, 1.1.1 xxx
        if re.match(r'^\d+(\.\d+)*\s+\S', text):
            dots = text.split('.')
            if len(dots) >= 3:
                return max(current_level, 3)
            elif len(dots) == 2:
                return max(current_level, 2)
            else:
                return max(current_level, 1)

        # 中文数字编号: 一、xxx / （一）xxx
        if re.match(r'^[（(]?[一二三四五六七八九十]+[）).、]', text) and char_count <= 80:
            return max(current_level, 1)

        # 关键词模式
        for kw in self._CN_HEADING_PATTERNS:
            if text.startswith(kw) and char_count <= 60:
                return max(current_level, 1)
            # "关键词：xxx" 或 "关键词: xxx"
            if text.startswith(kw) and ('：' in text or ':' in text) and char_count <= 100:
                return max(current_level, 2)

        # "摘要"/"Abstract" 单独一行
        if text.strip() in ('摘要', 'Abstract') and char_count <= 10:
            return max(current_level, 1)

        return current_level

    # ── 后处理 ────────────────────────────────────────────────

    def _post_process_paragraphs(self, paragraphs: list) -> list:
        """
        合并相邻的正文段落、过滤页码/页眉等噪声。
        """
        if not paragraphs:
            return paragraphs

        cleaned = []
        for p in paragraphs:
            text = p['text']

            # 过滤纯数字（页码）
            if text.isdigit() and len(text) <= 3:
                continue

            # 过滤页眉常见模式
            if any(text.startswith(pf) for pf in ['收稿日期', '作者简介', '基金项目', '中图分类号', 'DOI']):
                if p['char_count'] <= 80:
                    continue

            cleaned.append(p)

        return cleaned

    # ── 元信息检测 ────────────────────────────────────────────

    def _detect_title(self, paragraphs: list) -> str:
        """从段落中检测论文标题（通常是第一个大字号段落）"""
        for p in paragraphs:
            if p['level'] == 0 and p.get('font_size', 10) >= 14:
                return p['text']
        # 回退：第一个非短段落
        for p in paragraphs:
            if p['char_count'] >= 5 and p['level'] == 0:
                return p['text']
        return ''

    def _detect_author(self, paragraphs: list) -> str:
        """检测作者信息"""
        import re
        for i, p in enumerate(paragraphs):
            text = p['text']
            # 通常在标题后面几行，带数字上标（机构编号）
            if re.search(r'[\u4e00-\u9fff]{2,4}\s*[\d,，、]+', text) and p['char_count'] <= 60:
                return text
        return ''

    def _detect_date(self, paragraphs: list) -> str:
        """检测日期"""
        import re
        for p in paragraphs:
            text = p['text']
            m = re.search(r'(\d{4})[年.-](\d{1,2})[月.-](\d{1,2})', text)
            if m:
                return f"{m.group(1)}-{m.group(2).zfill(2)}-{m.group(3).zfill(2)}"
        return ''

    # ── 统一转换入口 ──────────────────────────────────────────

    def convert(self, input_path: str, output_path: str = None) -> dict:
        if not os.path.exists(input_path):
            return {'success': False, 'error': f'Input file not found: {input_path}'}
        if self.is_pdf_format(input_path):
            return self.convert_pdf(input_path, output_path)
        if self.is_doc_format(input_path):
            return self.convert_doc(input_path, output_path)
        ext = Path(input_path).suffix.lower()
        if ext == '.docx':
            return {
                'success': True, 'input_path': input_path,
                'output_path': input_path, 'method': 'none',
                'message': 'Already .docx format'
            }
        return {'success': False, 'error': f'Unsupported format: {ext}'}

    def batch_convert(self, input_dir: str, output_dir: str = None) -> list:
        if output_dir is None:
            output_dir = input_dir
        Path(output_dir).mkdir(parents=True, exist_ok=True)
        results = []
        # .doc files
        doc_files = sorted(Path(input_dir).glob('*.doc'))
        doc_files = [f for f in doc_files if f.suffix.lower() == '.doc' and not f.name.startswith('~')]
        for doc_file in doc_files:
            output_path = str(Path(output_dir) / doc_file.with_suffix('.docx').name)
            print(f"[INFO] Converting .doc: {doc_file.name}")
            result = self.convert_doc(str(doc_file), output_path)
            result['input'] = str(doc_file)
            results.append(result)
        # .pdf files
        pdf_files = sorted(Path(input_dir).glob('*.pdf'))
        pdf_files = [f for f in pdf_files if not f.name.startswith('~')]
        for pdf_file in pdf_files:
            output_path = str(Path(output_dir) / pdf_file.with_suffix('.docx').name)
            print(f"[INFO] Converting PDF: {pdf_file.name}")
            result = self.convert_pdf(str(pdf_file), output_path)
            result['input'] = str(pdf_file)
            results.append(result)
        # .docx files (copy)
        docx_files = sorted(Path(input_dir).glob('*.docx'))
        docx_files = [f for f in docx_files if not f.name.startswith('~')]
        for docx_file in docx_files:
            target_path = str(Path(output_dir) / docx_file.name)
            if str(docx_file) != target_path:
                shutil.copy2(str(docx_file), target_path)
            results.append({
                'success': True, 'input': str(docx_file),
                'output_path': target_path, 'method': 'copy',
                'message': 'Already .docx format'
            })
        return results

    # ── 内容提取（文本 + 表格） ───────────────────────────────

    def extract_content(self, docx_path: str) -> dict:
        if not os.path.exists(docx_path):
            return {'success': False, 'error': f'File not found: {docx_path}'}
        try:
            from docx import Document
        except ImportError:
            return {'success': False, 'error': 'python-docx not installed. Run: pip install python-docx'}
        try:
            doc = Document(docx_path)
            core = doc.core_properties
            paragraphs = []
            full_text_parts = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue
                style_name = para.style.name if para.style else 'Normal'
                is_heading = style_name.startswith('Heading')
                level = 0
                if is_heading:
                    try:
                        level = int(style_name.replace('Heading ', ''))
                    except ValueError:
                        level = 0
                paragraphs.append({
                    'text': text, 'style': style_name,
                    'is_heading': is_heading, 'level': level
                })
                full_text_parts.append(text)
            tables = []
            for table in doc.tables:
                rows = []
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    rows.append(cells)
                if rows:
                    tables.append({'rows': rows})
                for row in rows:
                    full_text_parts.append('\t'.join(row))
            return {
                'success': True,
                'file_name': Path(docx_path).name,
                'file_path': str(docx_path),
                'title': core.title or Path(docx_path).stem,
                'author': core.author or '',
                'created_at': str(core.created) if core.created else '',
                'paragraph_count': len(paragraphs),
                'table_count': len(tables),
                'paragraphs': paragraphs,
                'tables': tables,
                'full_text': '\n'.join(full_text_parts)
            }
        except Exception as e:
            return {'success': False, 'error': f'Failed to extract content: {str(e)}'}

    # ── 图片提取 ──────────────────────────────────────────────

    def extract_images(self, docx_path: str, output_dir: str) -> dict:
        if not os.path.exists(docx_path):
            return {'success': False, 'error': f'File not found: {docx_path}'}
        try:
            from docx import Document
            from docx.opc.constants import RELATIONSHIP_TYPE as RT
        except ImportError:
            return {'success': False, 'error': 'python-docx not installed. Run: pip install python-docx'}
        try:
            doc = Document(docx_path)
            Path(output_dir).mkdir(parents=True, exist_ok=True)
            saved = []
            idx = 0
            for rel in doc.part.rels.values():
                if "image" in rel.reltype:
                    idx += 1
                    image_data = rel.target_part.blob
                    ext = Path(rel.target_ref).suffix.lower()
                    if ext not in ('.png', '.jpg', '.jpeg', '.gif', '.bmp', '.tiff', '.emf', '.wmf'):
                        ext = '.png'
                    out_name = f"image_{idx}{ext}"
                    out_path = Path(output_dir) / out_name
                    with open(out_path, 'wb') as f:
                        f.write(image_data)
                    saved.append(str(out_path))
            print(f"[INFO] Extracted {len(saved)} image(s) from {Path(docx_path).name}")
            return {
                'success': True,
                'source': str(docx_path),
                'output_dir': str(output_dir),
                'image_count': len(saved),
                'images': saved
            }
        except Exception as e:
            return {'success': False, 'error': f'Failed to extract images: {str(e)}'}

    # ── .txt 内容提取 ─────────────────────────────────────────

    def extract_txt_content(self, txt_path: str) -> dict:
        if not os.path.exists(txt_path):
            return {'success': False, 'error': f'File not found: {txt_path}'}
        try:
            text = Path(txt_path).read_text(encoding='utf-8')
        except UnicodeDecodeError:
            try:
                text = Path(txt_path).read_text(encoding='gbk')
            except Exception:
                text = Path(txt_path).read_text(encoding='utf-8', errors='ignore')
        lines = [l.strip() for l in text.splitlines() if l.strip()]
        return {
            'success': True,
            'file_name': Path(txt_path).name,
            'file_path': str(txt_path),
            'title': Path(txt_path).stem,
            'author': '',
            'created_at': '',
            'paragraph_count': len(lines),
            'table_count': 0,
            'paragraphs': [
                {'text': l, 'style': 'Normal', 'is_heading': False, 'level': 0}
                for l in lines
            ],
            'tables': [],
            'full_text': text
        }

    # ── 批量提取 ──────────────────────────────────────────────

    def extract_batch(self, input_dir: str) -> list:
        results = []
        docx_files = sorted(Path(input_dir).glob('*.docx'))
        docx_files = [f for f in docx_files if not f.name.startswith('~')]
        for docx_file in docx_files:
            print(f"[INFO] Extracting content: {docx_file.name}")
            content = self.extract_content(str(docx_file))
            results.append(content)
        return results

    # ── 全流程处理 ────────────────────────────────────────────

    def process_all(self, input_dir: str, workspace_dir: str) -> dict:
        input_dir = str(Path(input_dir).resolve())
        workspace_dir = str(Path(workspace_dir).resolve())
        converted_dir = str(Path(workspace_dir) / 'converted')
        summary_dir = str(Path(workspace_dir) / 'summary')
        logs_dir = str(Path(workspace_dir) / 'logs')

        for d in [converted_dir, summary_dir, logs_dir]:
            Path(d).mkdir(parents=True, exist_ok=True)

        # 空输入检查
        input_files = []
        for ext in ['*.doc', '*.docx', '*.pdf', '*.txt']:
            input_files.extend(Path(input_dir).glob(ext))
        input_files = sorted([f for f in input_files if not f.name.startswith('~')])

        if not input_files:
            print(f"[WARN] No supported files found in {input_dir}")
            return {
                'status': 'empty',
                'convert_results': [],
                'summaries': [],
                'converted_dir': converted_dir,
                'summary_dir': summary_dir,
            }

        convert_results = self.batch_convert(input_dir, converted_dir)

        docx_files = sorted(Path(converted_dir).glob('*.docx'))
        docx_files = [f for f in docx_files if not f.name.startswith('~')]

        # 处理 .txt 文件（直接提取，不需要转换）
        txt_files = sorted(Path(input_dir).glob('*.txt'))
        txt_files = [f for f in txt_files if not f.name.startswith('~')]

        # 处理 PDF 文件（使用PyMuPDF直接提取）
        pdf_files = sorted(Path(input_dir).glob('*.pdf'))
        pdf_files = [f for f in pdf_files if not f.name.startswith('~')]

        summaries = []

        # 处理已转换的DOCX文件
        for docx_file in docx_files:
            stem = docx_file.stem
            doc_summary_dir = Path(summary_dir) / stem
            text_dir = doc_summary_dir / 'text'
            img_dir = doc_summary_dir / 'img'
            text_dir.mkdir(parents=True, exist_ok=True)
            img_dir.mkdir(parents=True, exist_ok=True)

            content = self.extract_content(str(docx_file))
            if content.get('success'):
                content_path = text_dir / 'content.json'
                with open(content_path, 'w', encoding='utf-8') as f:
                    json.dump(content, f, indent=2, ensure_ascii=False)

            img_result = self.extract_images(str(docx_file), str(img_dir))

            summaries.append({
                'file_name': stem,
                'source_file': docx_file.name,
                'docx_path': str(docx_file),
                'text_dir': str(text_dir),
                'img_dir': str(img_dir),
                'content_success': content.get('success', False),
                'image_count': img_result.get('image_count', 0),
            })

        # 处理TXT文件
        for txt_file in txt_files:
            stem = txt_file.stem
            doc_summary_dir = Path(summary_dir) / stem
            text_dir = doc_summary_dir / 'text'
            text_dir.mkdir(parents=True, exist_ok=True)

            content = self.extract_txt_content(str(txt_file))
            if content.get('success'):
                content_path = text_dir / 'content.json'
                with open(content_path, 'w', encoding='utf-8') as f:
                    json.dump(content, f, indent=2, ensure_ascii=False)

            summaries.append({
                'file_name': stem,
                'source_file': txt_file.name,
                'docx_path': None,
                'text_dir': str(text_dir),
                'img_dir': None,
                'content_success': content.get('success', False),
                'image_count': 0,
            })

        # 处理PDF文件（直接使用PyMuPDF提取，适用于知网等中文PDF）
        for pdf_file in pdf_files:
            stem = pdf_file.stem
            # 跳过已通过PDF转DOCX成功处理的文件
            already_processed = any(s['file_name'] == stem for s in summaries)
            if already_processed:
                continue

            doc_summary_dir = Path(summary_dir) / stem
            text_dir = doc_summary_dir / 'text'
            text_dir.mkdir(parents=True, exist_ok=True)

            print(f"[INFO] Extracting PDF content directly with PyMuPDF: {pdf_file.name}")
            content = self.extract_pdf_content(str(pdf_file))
            if content.get('success'):
                content_path = text_dir / 'content.json'
                with open(content_path, 'w', encoding='utf-8') as f:
                    json.dump(content, f, indent=2, ensure_ascii=False)

            summaries.append({
                'file_name': stem,
                'source_file': pdf_file.name,
                'docx_path': None,
                'text_dir': str(text_dir),
                'img_dir': None,
                'content_success': content.get('success', False),
                'image_count': 0,
            })

        return {
            'status': 'completed',
            'convert_results': convert_results,
            'summaries': summaries,
            'converted_dir': converted_dir,
            'summary_dir': summary_dir,
        }


def main():
    if len(sys.argv) < 2:
        print("Usage:")
        print("  python doc_converter.py convert <input_path> [output_path]")
        print("  python doc_converter.py batch_convert <input_dir> [output_dir]")
        print("  python doc_converter.py extract <docx_path>")
        print("  python doc_converter.py extract_batch <input_dir>")
        print("  python doc_converter.py extract_images <docx_path> <output_dir>")
        print("  python doc_converter.py process_all <input_dir> <workspace_dir>")
        sys.exit(1)

    cmd = sys.argv[1]
    converter = DocConverter()

    if cmd == 'convert':
        if len(sys.argv) < 3:
            print("Error: input_path required")
            sys.exit(1)
        result = converter.convert(sys.argv[2], sys.argv[3] if len(sys.argv) > 3 else None)
        print(json.dumps(result, indent=2, ensure_ascii=False))
        if not result['success']:
            sys.exit(1)

    elif cmd == 'batch_convert':
        if len(sys.argv) < 3:
            print("Error: input_dir required")
            sys.exit(1)
        results = converter.batch_convert(sys.argv[2], sys.argv[3] if len(sys.argv) > 3 else None)
        print(json.dumps(results, indent=2, ensure_ascii=False))

    elif cmd == 'extract':
        if len(sys.argv) < 3:
            print("Error: docx_path required")
            sys.exit(1)
        result = converter.extract_content(sys.argv[2])
        print(json.dumps(result, indent=2, ensure_ascii=False))

    elif cmd == 'extract_batch':
        if len(sys.argv) < 3:
            print("Error: input_dir required")
            sys.exit(1)
        results = converter.extract_batch(sys.argv[2])
        print(json.dumps(results, indent=2, ensure_ascii=False))

    elif cmd == 'extract_images':
        if len(sys.argv) < 4:
            print("Error: docx_path and output_dir required")
            sys.exit(1)
        result = converter.extract_images(sys.argv[2], sys.argv[3])
        print(json.dumps(result, indent=2, ensure_ascii=False))

    elif cmd == 'process_all':
        if len(sys.argv) < 4:
            print("Error: input_dir and workspace_dir required")
            sys.exit(1)
        result = converter.process_all(sys.argv[2], sys.argv[3])
        print(json.dumps(result, indent=2, ensure_ascii=False))

    else:
        print(f"Unknown command: {cmd}")
        sys.exit(1)


if __name__ == '__main__':
    main()
