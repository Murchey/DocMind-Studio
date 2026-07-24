"""
Font Fixer — DOCX 正文字体强制修正器

功能：
  针对 pandoc 生成的 DOCX（默认 Calibri），强制所有段落的正文字体为宋体。

用法：
  python font_fixer.py <input.docx> [output.docx]
  python font_fixer.py create-ref <output.docx>   # 创建 pandoc 引用模板

检测/修正规则：
  1. 修改 DOCX 默认样式（styles.xml 中的 Normal/DefaultParagraphFont）
  2. 遍历所有段落，未显式设置字体的 run 统一使用宋体
  3. 遍历所有段落，含中文的 run 强制设为宋体，英文设为 Times New Roman
  4. 标题（Heading 样式）保持黑体不变
"""

import json, os, shutil, sys, zipfile, tempfile
from pathlib import Path
from xml.etree import ElementTree as ET
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ============================================================
# 字体常量
# ============================================================

BODY_FONT_CN = "宋体"           # 中文正文
BODY_FONT_EN = "Times New Roman"  # 英文正文
HEADING_FONT = "黑体"           # 标题

# 应保持不变的字体（不强制替换）
PRESERVED_FONTS = {"Consolas", "Courier New", "Monaco"}


# ============================================================
# Pandoc 引用模板生成
# ============================================================

def create_pandoc_reference(output_path: str):
    """
    创建一个最小化的 DOCX，设置默认字体为宋体，
    用作 pandoc --reference-doc 的模板。
    """
    doc = Document()
    
    # 1. 修改默认段落样式的字体
    style = doc.styles["Normal"]
    font = style.font
    font.name = BODY_FONT_EN
    font.size = Pt(12)
    
    # 2. 设置亚洲字体
    rpr = style.element.get_or_add_rPr()
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rpr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), BODY_FONT_EN)
    rFonts.set(qn("w:hAnsi"), BODY_FONT_EN)
    rFonts.set(qn("w:eastAsia"), BODY_FONT_CN)
    
    # 3. 设置默认段落格式
    pPr = style.element.get_or_add_pPr()
    spacing = pPr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        pPr.append(spacing)
    spacing.set(qn("w:line"), "360")  # 1.5 倍行距
    spacing.set(qn("w:lineRule"), "auto")
    
    # 4. 设置文档默认字体（DocDefaults）
    body = doc.element.body
    styles_elem = doc.styles.element
    
    # 5. 设置标题样式字体
    for level in range(1, 4):
        style_name = f"Heading {level}"
        try:
            hs = doc.styles[style_name]
            hfont = hs.font
            hfont.name = HEADING_FONT
            hfont.size = Pt(16 if level == 1 else 14 if level == 2 else 13)
            hfont.bold = True
            h_rpr = hs.element.get_or_add_rPr()
            h_rFonts = h_rpr.find(qn("w:rFonts"))
            if h_rFonts is None:
                h_rFonts = OxmlElement("w:rFonts")
                h_rpr.insert(0, h_rFonts)
            h_rFonts.set(qn("w:ascii"), HEADING_FONT)
            h_rFonts.set(qn("w:hAnsi"), HEADING_FONT)
            h_rFonts.set(qn("w:eastAsia"), HEADING_FONT)
        except KeyError:
            pass
    
    # 6. 保存
    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    
    return {"status": "ok", "path": output_path, "body_font_cn": BODY_FONT_CN, "body_font_en": BODY_FONT_EN}


# ============================================================
# DOCX 字体修正引擎
# ============================================================

def fix_docx_fonts(input_path: str, output_path: str = None) -> dict:
    """
    强制修正 DOCX 文件的所有正文字体为宋体。

    Args:
        input_path: 输入 DOCX 路径
        output_path: 输出 DOCX 路径（不指定则覆盖原文件）

    Returns:
        dict: {"status": "ok", "fixed_runs": N, "fixed_styles": N}
    """
    input_path = Path(input_path)
    if not input_path.exists():
        return {"status": "error", "message": f"文件不存在: {input_path}"}
    
    if output_path is None:
        output_path = str(input_path)
    
    doc = Document(str(input_path))
    fixed_runs = 0
    fixed_styles = 0
    
    # --- 1. 修正默认样式 ---
    try:
        normal_style = doc.styles["Normal"]
        rpr = normal_style.element.get_or_add_rPr()
        rFonts = rpr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rpr.insert(0, rFonts)
        else:
            old_ascii = rFonts.get(qn("w:ascii"))
            old_east = rFonts.get(qn("w:eastAsia"))
            if old_ascii != BODY_FONT_EN:
                fixed_styles += 1
            if old_east != BODY_FONT_CN:
                fixed_styles += 1
        rFonts.set(qn("w:ascii"), BODY_FONT_EN)
        rFonts.set(qn("w:hAnsi"), BODY_FONT_EN)
        rFonts.set(qn("w:eastAsia"), BODY_FONT_CN)
        normal_style.font.name = BODY_FONT_EN
        normal_style.font.size = Pt(12)
    except KeyError:
        pass
    
    # --- 2. 修正每个段落的字体 ---
    for para in doc.paragraphs:
        style_name = para.style.name if para.style else ""
        is_heading = style_name.lower().startswith("heading")
        
        for run in para.runs:
            text = run.text.strip()
            if not text:
                continue
            
            has_chinese = any("\u4e00" <= c <= "\u9fff" for c in text)
            current_font = run.font.name or ""
            
            # 检查是否是应保留的字体
            if current_font in PRESERVED_FONTS:
                continue
            
            # 设置正确的字体
            if is_heading:
                target_font = HEADING_FONT
            else:
                target_font = BODY_FONT_CN if has_chinese else BODY_FONT_EN
            
            run.font.name = target_font
            rpr = run._element.get_or_add_rPr()
            rFonts = rpr.find(qn("w:rFonts"))
            if rFonts is None:
                rFonts = OxmlElement("w:rFonts")
                rpr.insert(0, rFonts)
            rFonts.set(qn("w:ascii"), target_font)
            rFonts.set(qn("w:hAnsi"), target_font)
            if has_chinese:
                rFonts.set(qn("w:eastAsia"), BODY_FONT_CN)
            elif is_heading:
                rFonts.set(qn("w:eastAsia"), HEADING_FONT)
            
            fixed_runs += 1
    
    # --- 3. 修正表格中的字体 ---
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        text = run.text.strip()
                        if not text:
                            continue
                        has_chinese = any("\u4e00" <= c <= "\u9fff" for c in text)
                        if has_chinese:
                            run.font.name = BODY_FONT_CN
                            rpr = run._element.get_or_add_rPr()
                            rFonts = rpr.find(qn("w:rFonts"))
                            if rFonts is None:
                                rFonts = OxmlElement("w:rFonts")
                                rpr.insert(0, rFonts)
                            rFonts.set(qn("w:ascii"), BODY_FONT_CN)
                            rFonts.set(qn("w:hAnsi"), BODY_FONT_CN)
                            rFonts.set(qn("w:eastAsia"), BODY_FONT_CN)
                            fixed_runs += 1
    
    # --- 4. 保存 ---
    if output_path != input_path:
        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    
    doc.save(str(output_path))
    
    return {
        "status": "ok",
        "input_path": str(input_path),
        "output_path": str(output_path),
        "fixed_runs": fixed_runs,
        "fixed_styles": fixed_styles,
        "body_font_cn": BODY_FONT_CN,
        "body_font_en": BODY_FONT_EN,
        "heading_font": HEADING_FONT,
    }


# ============================================================
# CLI
# ============================================================

def main():
    if len(sys.argv) < 2:
        print("Font Fixer — DOCX 正文字体强制修正器")
        print("")
        print("用法:")
        print("  python font_fixer.py fix <input.docx> [output.docx]    # 修正字体")
        print("  python font_fixer.py create-ref <output.docx>          # 创建 pandoc 引用模板")
        print("")
        print(f"正文字体: {BODY_FONT_CN} / {BODY_FONT_EN}")
        print(f"标题字体: {HEADING_FONT}")
        sys.exit(0)
    
    command = sys.argv[1]
    
    if command == "fix":
        if len(sys.argv) < 3:
            print("请指定 DOCX 文件路径")
            sys.exit(1)
        result = fix_docx_fonts(sys.argv[2], sys.argv[3] if len(sys.argv) > 3 else None)
        if result["status"] == "error":
            print(f"[ERROR] {result['message']}")
            sys.exit(1)
        result["_message"] = f"已修正 {result['fixed_runs']} 个 run, {result['fixed_styles']} 个样式"
        print(json.dumps(result, ensure_ascii=False, indent=2))
    
    elif command == "create-ref":
        if len(sys.argv) < 3:
            print("请指定输出路径")
            sys.exit(1)
        result = create_pandoc_reference(sys.argv[2])
        print(json.dumps(result, ensure_ascii=False, indent=2))
        print(f"[OK] pandoc 引用模板已创建: {result['path']}")
    
    else:
        print(f"未知命令: {command}")
        sys.exit(1)


if __name__ == "__main__":
    main()
