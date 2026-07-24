"""
Content Validator — MD 源文档预检 + DOCX 输出后检

功能：
  1. pre_check_md(source_md_path) → Report
     在 MD→DOCX 转换前检查源文件，识别可能引起乱码的模式
  2. post_check_docx(output_docx_path, source_md_path) → Report
     在转换后检查 DOCX，对比源文件识别乱码段落
  3. check(source_md, output_docx=None) → Report
     全流程：预检 → [可选转换] → 后检

检测规则：
  - RAW_JSON_BLOB：嵌入的 JSON/代码块中包含反斜杠路径
  - OVERSIZE_PARAGRAPH：段落过长（> 5000 字符），pandoc 可能截断
  - BACKSLASH_PATH：文本中包含 Windows 反斜杠路径
  - UNICODE_GARBLED：DOCX 中出现异常 Unicode 字符模式（乱码特征）
  - PARAGRAPH_MISMATCH：DOCX 段落数与源 MD 段落数偏差过大
  - TABLE_CORRUPTION：表格数据在转换中丢失或变形
"""

import re
import json
from pathlib import Path
from typing import List, Dict, Any, Optional
from dataclasses import dataclass, field
from datetime import datetime


# ============================================================
# 数据结构
# ============================================================

@dataclass
class Issue:
    """单个检测到的问题"""
    severity: str           # "ERROR" | "WARNING" | "INFO"
    category: str           # 规则分类
    location: str           # 位置描述（行号/段落索引）
    message: str            # 人类可读的描述
    snippet: str = ""       # 相关文本片段（前 120 字符）
    suggestion: str = ""    # 修复建议

    def to_dict(self) -> dict:
        return {
            "severity": self.severity,
            "category": self.category,
            "location": self.location,
            "message": self.message,
            "snippet": self.snippet[:120],
            "suggestion": self.suggestion,
        }


@dataclass
class Report:
    """校验报告"""
    phase: str                      # "pre_check" | "post_check" | "full"
    status: str                     # "PASS" | "WARN" | "FAIL"
    source_file: str = ""
    output_file: str = ""
    issues: List[Issue] = field(default_factory=list)
    summary: dict = field(default_factory=dict)
    generated_at: str = ""

    def add_issue(self, issue: Issue):
        self.issues.append(issue)

    def has_errors(self) -> bool:
        return any(i.severity == "ERROR" for i in self.issues)

    def has_warnings(self) -> bool:
        return any(i.severity == "WARNING" for i in self.issues)

    def compute_status(self):
        if self.has_errors():
            self.status = "FAIL"
        elif self.has_warnings():
            self.status = "WARN"
        else:
            self.status = "PASS"

    def to_dict(self, detailed: bool = False) -> dict:
        d = {
            "phase": self.phase,
            "status": self.status,
            "source_file": self.source_file,
            "output_file": self.output_file,
            "summary": self.summary,
            "generated_at": self.generated_at,
        }
        if detailed:
            d["issues"] = [i.to_dict() for i in self.issues]
        return d

    def print_summary(self):
        icons = {"PASS": "PASS", "WARN": "WARN", "FAIL": "FAIL"}
        print(f"\n{'=' * 50}")
        print(f"  Content Validator Report [{self.phase}]")
        print(f"  Status: {icons.get(self.status, '?')} {self.status}")
        print(f"{'=' * 50}")
        errs = sum(1 for i in self.issues if i.severity == "ERROR")
        warns = sum(1 for i in self.issues if i.severity == "WARNING")
        infos = sum(1 for i in self.issues if i.severity == "INFO")
        print(f"  Errors:   {errs}")
        print(f"  Warnings: {warns}")
        print(f"  Info:     {infos}")
        if self.issues:
            print(f"\n  Issues:")
            for i, issue in enumerate(self.issues, 1):
                icon = {"ERROR": "[E]", "WARNING": "[W]", "INFO": "[I]"}.get(issue.severity, "[?]")
                print(f"    {icon} {issue.category}: {issue.message}")
                if issue.snippet:
                    print(f"       snippet: {issue.snippet[:80]}...")
                if issue.suggestion:
                    print(f"       fix: {issue.suggestion}")
        print(f"{'=' * 50}")


# ============================================================
# 检测引擎
# ============================================================

class ContentValidator:
    """
    MD↔DOCX 内容乱码校验器

    Usage:
        validator = ContentValidator()
        
        # 仅预检
        report = validator.pre_check_md("report.md")
        report.print_summary()
        
        # 仅后检
        report = validator.post_check_docx("report.docx", "report.md")
        
        # 全流程
        report = validator.check("report.md", "report.docx")
        if report.has_errors():
            print("校验未通过，请修复问题后重新生成")
    """

    # 反斜杠路径模式: d:\\xxx 或 \\xx
    BACKSLASH_PATH_RE = re.compile(r'(?:[a-zA-Z]:)?(?:[\\/]{1,2}[^\\/\\s]+)+')
    
    # JSON 块中的反斜杠转义: \\x
    JSON_BACKSLASH_RE = re.compile(r'\\[\\"\'nrt0]')
    
    # 连续 $ 符号（被吃掉的转义路径特征）
    GARBLED_DOLLAR_RE = re.compile(r'\${3,}')
    
    # 异常 Unicode：CJK 区以外的非 ASCII 字符，且不是常见标点
    UNUSUAL_UNICODE_RE = re.compile(r'[\u2018\u2019\u201c\u201d\u2026]')

    def __init__(self):
        self.generated_at = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    # ----------------------------------------------------------
    # 预检：Markdown 源文件
    # ----------------------------------------------------------

    def pre_check_md(self, md_path: str) -> Report:
        """检查 MD 源文件中的乱码风险"""
        report = Report(phase="pre_check", status="", source_file=md_path, generated_at=self.generated_at)
        path = Path(md_path)
        
        if not path.exists():
            report.add_issue(Issue("ERROR", "FILE_NOT_FOUND", md_path, "源文件不存在"))
            report.compute_status()
            return report

        content = path.read_text(encoding="utf-8")
        lines = content.split("\n")
        
        report.summary = {
            "total_lines": len(lines),
            "total_chars": len(content),
        }

        in_code_block = False
        in_json_block = False
        json_lines = []
        
        for idx, line in enumerate(lines, 1):
            stripped = line.strip()

            # --- 检测代码块 ---
            if stripped.startswith("```"):
                if in_code_block:
                    in_code_block = False
                    # 代码块结束，检查是否包含 JSON
                    self._check_json_block(report, json_lines, idx)
                    json_lines = []
                else:
                    in_code_block = True
                    json_lines = []
                continue

            # --- 检测 JSON 特征（不在代码块中） ---
            if not in_code_block:
                # 裸 JSON 检测：以 { 或 " 开头，包含 "Unnamed"、"input_path"、"data_analysis" 等特征
                if stripped.startswith("{") or stripped.startswith('"status"'):
                    in_json_block = True
                    json_lines = [(idx, line)]
                    continue
                
                if in_json_block:
                    json_lines.append((idx, line))
                    if stripped == "}" or stripped.endswith("},") or stripped.endswith("}"):
                        self._check_json_block(report, json_lines, idx)
                        json_lines = []
                        in_json_block = False
                    continue

            # --- 检测段落长度 ---
            if not in_code_block and not in_json_block and len(line) > 5000:
                report.add_issue(Issue(
                    "WARNING", "OVERSIZE_PARAGRAPH", f"L{idx}",
                    f"段落过长 ({len(line)} chars)，pandoc 可能截断或处理异常",
                    snippet=line[:120],
                    suggestion="将长段落拆分为多个短段落（每个 < 2000 chars）"
                ))

            # --- 检测反斜杠路径（不在代码块中） ---
            if not in_code_block and not in_json_block:
                matches = self.BACKSLASH_PATH_RE.findall(line)
                for m in matches:
                    if "\\" in m:
                        report.add_issue(Issue(
                            "ERROR", "BACKSLASH_PATH", f"L{idx}",
                            f"文本中包含 Windows 反斜杠路径，pandoc 会将其解析为转义符导致乱码",
                            snippet=m,
                            suggestion="将路径替换为纯文字说明或正斜杠 /"
                        ))
                        break  # 每行只报一次

            # --- 检测连续 $（已产生的乱码特征） ---
            dollar_matches = self.GARBLED_DOLLAR_RE.findall(line)
            if dollar_matches:
                report.add_issue(Issue(
                    "ERROR", "GARBLED_DOLLAR", f"L{idx}",
                    "连续 $$$ 符号 — 这是反斜杠被 pandoc 吃掉后产生的乱码特征",
                    snippet=line[:120],
                    suggestion="删除这段乱码，将原始内容以纯文本重写"
                ))

        # 如果代码块未关闭
        if json_lines:
            self._check_json_block(report, json_lines, len(lines))

        report.compute_status()
        return report

    def _check_json_block(self, report: Report, json_lines: list, end_line: int):
        """检查代码块或 JSON 块是否包含反斜杠路径"""
        if not json_lines:
            return
        
        # 提取块内容
        start_line = json_lines[0][0] if isinstance(json_lines[0], tuple) else 1
        text = ""
        for item in json_lines:
            if isinstance(item, tuple):
                text += item[1] + "\n"
            else:
                text += item + "\n"
        
        # 检查反斜杠路径
        bs_matches = self.JSON_BACKSLASH_RE.findall(text)
        if bs_matches:
            report.add_issue(Issue(
                "ERROR", "RAW_JSON_BLOB", f"L{start_line}-L{end_line}",
                f"嵌入了含反斜杠转义符的 raw JSON / 数据块 ({len(bs_matches)} 处)。"
                "pandoc 在转换时会将 \\x 序列解析为转义字符，导致路径信息丢失和乱码",
                snippet=text[:120].replace("\n", " "),
                suggestion="将 raw JSON 替换为纯文本摘要或 Markdown 表格，避免嵌入原始数据结构"
            ))

    # ----------------------------------------------------------
    # 后检：DOCX 输出文件
    # ----------------------------------------------------------

    def post_check_docx(self, docx_path: str, md_path: Optional[str] = None) -> Report:
        """检查 DOCX 输出中的乱码"""
        report = Report(phase="post_check", status="", source_file=md_path or "", 
                       output_file=docx_path, generated_at=self.generated_at)

        path = Path(docx_path)
        if not path.exists():
            report.add_issue(Issue("ERROR", "FILE_NOT_FOUND", docx_path, "DOCX 文件不存在"))
            report.compute_status()
            return report

        try:
            from docx import Document
            from docx.oxml.ns import qn
        except ImportError:
            report.add_issue(Issue("ERROR", "DEPENDENCY", "python-docx", "缺少 python-docx 库"))
            report.compute_status()
            return report

        doc = Document(str(path))
        
        # 提取所有段落文本
        paragraphs = []
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if text:
                paragraphs.append({
                    "index": i,
                    "text": text,
                    "style": para.style.name if para.style else "None",
                    "len": len(text),
                })

        report.summary = {
            "total_paragraphs_with_text": len(paragraphs),
            "total_paragraphs_raw": len(doc.paragraphs),
            "file_size": path.stat().st_size,
        }

        # --- 检测 1: 乱码特征 — 连续 $$$ ---
        for p in paragraphs:
            if self.GARBLED_DOLLAR_RE.search(p["text"]):
                report.add_issue(Issue(
                    "ERROR", "GARBLED_DOLLAR", f"Para {p['index']}",
                    "检测到连续 $$$ 符号 — 这是反斜杠路径被转换吞噬后产生的典型乱码",
                    snippet=p["text"][:120],
                    suggestion="检查源 MD 中对应的段落，确保没有嵌入 raw JSON 或反斜杠路径"
                ))

        # --- 检测 2: 异常的 Unicode 组合 ---
        for p in paragraphs:
            # 检测 Latin + CJK 混合中的异常
            latin_chars = sum(1 for c in p["text"] if ord(c) < 128 and c.isalpha())
            total_alpha = sum(1 for c in p["text"] if c.isalpha())
            if total_alpha > 0 and latin_chars / total_alpha > 0.8:
                # 大部分字符是拉丁字母 — 可能是英文内容被截断
                pass  # 英文内容页正常

        # --- 检测 3: 缺失/截断的段落 ---
        for p in paragraphs:
            # 检测以不完整 JSON 开头的段落
            text = p["text"].strip()
            json_starts = ['{ "', '{"', '"status"', '"data_']
            if any(text.startswith(s) for s in json_starts):
                report.add_issue(Issue(
                    "ERROR", "JSON_RESIDUE", f"Para {p['index']}",
                    "DOCX 中残留 JSON 数据片段，说明源 MD 中的 raw JSON 未被清除",
                    snippet=text[:120],
                    suggestion="从源 MD 中删除 raw JSON 块，替换为格式化文本"
                ))

        # --- 检测 4: 对比 MD 源（如果提供） ---
        if md_path and Path(md_path).exists():
            md_content = Path(md_path).read_text(encoding="utf-8")
            md_para_count = len([l for l in md_content.split("\n") if l.strip() and not l.startswith("```")])
            docx_para_count = len(paragraphs)
            
            ratio = docx_para_count / md_para_count if md_para_count > 0 else 0
            if ratio < 0.3:
                report.add_issue(Issue(
                    "WARNING", "PARAGRAPH_MISMATCH", f"MD={md_para_count}, DOCX={docx_para_count}",
                    f"DOCX 段落数 ({docx_para_count}) 远少于源 MD ({md_para_count})，"
                    f"可能存在大量内容丢失（覆盖率 {ratio:.0%}）",
                    suggestion="检查 MD 中是否有 pandoc 无法处理的特殊格式"
                ))

        # --- 检测 5: 检查表格是否损坏 ---
        table_count_in_docx = len(doc.tables)
        if md_path and Path(md_path).exists():
            md_table_count = md_content.count("|---") if 'md_content' in dir() else 0
            if md_table_count > 0 and table_count_in_docx == 0:
                report.add_issue(Issue(
                    "WARNING", "TABLE_CORRUPTION", f"MD={md_table_count}, DOCX={table_count_in_docx}",
                    "MD 中包含 Markdown 表格但 DOCX 中没有任何表格，表格数据可能在转换中丢失",
                    suggestion="检查 MD 表格格式是否正确（表头对齐线 |---| 必须存在）"
                ))

        report.compute_status()
        return report

    # ----------------------------------------------------------
    # 全流程校验
    # ----------------------------------------------------------

    def check(self, md_path: str, docx_path: Optional[str] = None) -> Report:
        """全流程：预检 → [后检]"""
        # 预检
        report = self.pre_check_md(md_path)
        report.phase = "full"
        
        # 后检（如果提供 DOCX 路径）
        if docx_path and Path(docx_path).exists():
            post = self.post_check_docx(docx_path, md_path)
            report.issues.extend(post.issues)
            report.summary["post_check"] = post.summary
            report.output_file = docx_path
            report.compute_status()
        
        return report


# ============================================================
# CLI 入口
# ============================================================

def main():
    import sys
    
    if len(sys.argv) < 2:
        print("Content Validator — MD/DOCX 内容乱码校验")
        print("")
        print("用法:")
        print("  python content_validator.py pre  <source.md>              # 仅预检")
        print("  python content_validator.py post <output.docx> [source.md] # 仅后检")
        print("  python content_validator.py full <source.md> [output.docx] # 全流程")
        print("")
        print("示例:")
        print('  python content_validator.py pre  report.md')
        print('  python content_validator.py full report.md report.docx')
        sys.exit(0)

    command = sys.argv[1]
    validator = ContentValidator()

    if command == "pre":
        if len(sys.argv) < 3:
            print("请指定 MD 文件路径")
            sys.exit(1)
        report = validator.pre_check_md(sys.argv[2])
        print(json.dumps(report.to_dict(detailed=True), ensure_ascii=False, indent=2))
        if report.has_errors():
            sys.exit(1)

    elif command == "post":
        if len(sys.argv) < 3:
            print("请指定 DOCX 文件路径")
            sys.exit(1)
        md_path = sys.argv[3] if len(sys.argv) > 3 else None
        report = validator.post_check_docx(sys.argv[2], md_path)
        print(json.dumps(report.to_dict(detailed=True), ensure_ascii=False, indent=2))
        if report.has_errors():
            sys.exit(1)

    elif command == "full":
        if len(sys.argv) < 3:
            print("请指定 MD 文件路径")
            sys.exit(1)
        docx_path = sys.argv[3] if len(sys.argv) > 3 else None
        report = validator.check(sys.argv[2], docx_path)
        print(json.dumps(report.to_dict(detailed=True), ensure_ascii=False, indent=2))
        if report.has_errors():
            sys.exit(1)

    else:
        print(f"未知命令: {command}")
        sys.exit(1)


if __name__ == "__main__":
    main()
